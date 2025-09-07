"""
Attachment Router: Handle XLS/XLSX/CSV/DOCX/PDF/Images
"""

import io
from typing import Dict, List, Optional
from email.message import EmailMessage
import logging
import pandas as pd
import pdfplumber
from docx import Document

from PIL import Image


class AttachmentData:
    """Container for attachment data"""
    
    def __init__(self, filename: str, content_type: str):
        self.filename = filename
        self.content_type = content_type
        self.structured_data: List[Dict] = [] 
        self.text_content: str = ""
        self.extraction_method: str = ""
        self.success: bool = False


class AttachmentRouter:
    """
    Route attachments to appropriate extractors
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        self.header_mappings = {
            'npi': ['npi', 'npi #', 'npi number', 'national provider identifier'],
            'tin': ['tin', 'tax id', 'federal id', 'ein', 'employer id'],
            'provider_name': ['provider name', 'provider', 'doctor name', 'physician'],
            'organization': ['organization', 'org', 'group', 'practice', 'medical group'],
            'specialty': ['specialty', 'speciality', 'practice specialty'],
            'license': ['license', 'state license', 'medical license', 'lic #'],
            'phone': ['phone', 'telephone', 'phone number', 'contact'],
            'fax': ['fax', 'fax number', 'facsimile'],
            'address': ['address', 'practice address', 'location'],
            'ppg': ['ppg', 'ppg id', 'practice group', 'group id'],
            'lob': ['lob', 'line of business', 'business line', 'network'],
        }
    
    def extract_attachments(self, msg: EmailMessage) -> List[AttachmentData]:
        """Extract data from all attachments"""
        attachments = []
        
        for part in msg.walk():
            if part.get_content_disposition() == 'attachment':
                filename = part.get_filename()
                if filename:
                    content_type = part.get_content_type()
                    payload = part.get_payload(decode=True)
                    
                    if payload:
                        attachment_data = self._route_attachment(filename, content_type, payload)
                        if attachment_data:
                            attachments.append(attachment_data)
        
        return attachments
    
    def _route_attachment(self, filename: str, content_type: str, payload: bytes) -> Optional[AttachmentData]:
        """Route attachment to appropriate extractor"""
        filename_lower = filename.lower()
        
        try:
            if filename_lower.endswith(('.xlsx', '.xls', '.csv')):
                return self._extract_spreadsheet(filename, content_type, payload)
            
            elif filename_lower.endswith('.docx'):
                return self._extract_docx(filename, content_type, payload)
            
            elif filename_lower.endswith('.pdf'):
                return self._extract_pdf(filename, content_type, payload)
            
            elif filename_lower.endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp')):
                return self._extract_image(filename, content_type, payload)
            
            else:
                self.logger.info(f"Unsupported attachment type: {filename}")
                return None
                
        except Exception as e:
            self.logger.error(f"Error processing attachment {filename}: {str(e)}")
            return None
    
    def _extract_spreadsheet(self, filename: str, content_type: str, payload: bytes) -> AttachmentData:
        """Extract data from XLS/XLSX/CSV files"""
        attachment = AttachmentData(filename, content_type)
        
        
        try:
            if filename.lower().endswith('.csv'):
                df = pd.read_csv(io.BytesIO(payload))
            else:
                df = pd.read_excel(io.BytesIO(payload))
            
            # Map headers using fuzzy matching
            mapped_data = self._map_spreadsheet_headers(df)
            attachment.structured_data = mapped_data
            attachment.extraction_method = "pandas_spreadsheet"
            attachment.success = True
            
            # Also create text representation
            attachment.text_content = df.to_string()
            
        except Exception as e:
            self.logger.error(f"Spreadsheet extraction failed for {filename}: {str(e)}")
            attachment.text_content = f"Error reading spreadsheet: {str(e)}"
        
        return attachment
    
    def _map_spreadsheet_headers(self, df) -> List[Dict]:
        """Map spreadsheet headers to standard fields using fuzzy matching"""
        if len(df) == 0:
            return []
        
        column_mapping = {}
        df_columns_lower = [col.lower().strip() for col in df.columns]
        
        for standard_field, variants in self.header_mappings.items():
            for col_idx, df_col in enumerate(df_columns_lower):
                for variant in variants:
                    if variant in df_col or self._fuzzy_match(variant, df_col):
                        column_mapping[df.columns[col_idx]] = standard_field
                        break
                if df.columns[col_idx] in column_mapping:
                    break
        
        structured_data = []
        for _, row in df.iterrows():
            row_data = {}
            for original_col, value in row.items():
                if original_col in column_mapping:
                    standard_field = column_mapping[original_col]
                    row_data[standard_field] = str(value) if pd.notna(value) else ""
            
            if row_data:
                structured_data.append(row_data)
        
        return structured_data
    
    def _fuzzy_match(self, pattern: str, text: str, threshold: float = 0.8) -> bool:
        """Simple fuzzy matching for header mapping"""
        pattern_words = set(pattern.split())
        text_words = set(text.split())
        
        if not pattern_words:
            return False
        
        intersection = pattern_words.intersection(text_words)
        return len(intersection) / len(pattern_words) >= threshold
    
    def _extract_docx(self, filename: str, content_type: str, payload: bytes) -> AttachmentData:
        """Extract tables and bullet lists from DOCX"""
        attachment = AttachmentData(filename, content_type)
        
       
        try:
            doc = Document(io.BytesIO(payload))
            
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            table_data = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = {}
                    cells = [cell.text.strip() for cell in row.cells]
                    
                    if len(cells) >= 2:
                        table_data.append(dict(zip(range(len(cells)), cells)))
            
            attachment.text_content = '\n'.join(full_text)
            attachment.structured_data = table_data if table_data else []
            attachment.extraction_method = "python_docx"
            attachment.success = True
            
        except Exception as e:
            self.logger.error(f"DOCX extraction failed for {filename}: {str(e)}")
            attachment.text_content = f"Error reading DOCX: {str(e)}"
        
        return attachment
    
    def _extract_pdf(self, filename: str, content_type: str, payload: bytes) -> AttachmentData:
        """Extract text and tables from PDF, with OCR fallback"""
        attachment = AttachmentData(filename, content_type)
        
       
        try:
            with pdfplumber.open(io.BytesIO(payload)) as pdf:
                full_text = []
                table_data = []
                
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text.append(text)
                    
                    tables = page.extract_tables()
                    for table in tables:
                        if table and len(table) > 0:
                            headers = table[0] if table[0] else []
                            for row in table[1:]:
                                if row:
                                    row_dict = dict(zip(headers, row))
                                    table_data.append(row_dict)
                
                attachment.text_content = '\n'.join(full_text)
                attachment.structured_data = table_data
                attachment.extraction_method = "pdfplumber"
                
                if not attachment.text_content.strip():
                    attachment = self._ocr_pdf(attachment, payload)
                else:
                    attachment.success = True
                    
        except Exception as e:
            self.logger.error(f"PDF extraction failed for {filename}: {str(e)}")
            attachment.text_content = f"Error reading PDF: {str(e)}"
            
            
        return attachment
    
    def _ocr_pdf(self, attachment: AttachmentData, payload: bytes) -> AttachmentData:
        """OCR PDF using Tesseract"""
        try:
            
            attachment.text_content = "OCR processing attempted but not fully implemented"
            attachment.extraction_method = "tesseract_ocr"
            attachment.success = False
            
        except Exception as e:
            attachment.text_content += f"\nOCR failed: {str(e)}"
        
        return attachment
    
    def _extract_image(self, filename: str, content_type: str, payload: bytes) -> AttachmentData:
        """Extract text from images using OCR"""
        attachment = AttachmentData(filename, content_type)
        
        
        try:
            image = Image.open(io.BytesIO(payload))
            
            if image.mode != 'L':
                image = image.convert('L')
            
            text = pytesseract.image_to_string(image)
            
            attachment.text_content = text
            attachment.extraction_method = "tesseract_image"
            attachment.success = bool(text.strip())
            
        except Exception as e:
            self.logger.error(f"Image OCR failed for {filename}: {str(e)}")
            attachment.text_content = f"Error processing image: {str(e)}"
        
        return attachment